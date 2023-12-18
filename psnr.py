import cv2
import os

import docx

from Matrix import StringMatrix
from docx import Document


# def get_ith_file_name(folder_path, i):
#     """
#     Get the name of the i-th file in the specified folder.
#
#     Parameters:
#     - folder_path (str): The path to the folder.
#     - i (int): The index of the file.
#
#     Returns:
#     - str: The name of the i-th file, or None if the index is out of range.
#     """
#     files = os.listdir(folder_path)
#     files = sorted(files)
#
#     if 0 <= i < len(files):
#         return files[i]
#     else:
#         print(f"Index {i} is out of range for the number of files in the folder.")
#         return None


def get_deionised_file_names(folder_path, original_image_name):
    target_image = original_image_name.split('.')[0]
    files = os.listdir(folder_path)

    files_names = []

    for file in files:
        if file.split("_")[0] == target_image:
            files_names.append(file)

    return files_names


def calculate_psnr(original_image, denoised_image):
    psnr = cv2.PSNR(original_image, denoised_image)
    return psnr


def main():
    original_folder = 'DS/Med'
    denoised_folder = 'DS_DeNoisyImages/Med'

    string_matrix = StringMatrix()

    for filename in os.listdir(original_folder):
        for denoised_image_path in get_deionised_file_names(denoised_folder, filename):
            if filename.endswith(('.png', '.jpg', '.jpeg', '.bmp')):
                original_path = os.path.join(original_folder, filename)
                denoised_path = os.path.join(denoised_folder, denoised_image_path)

                print("original image : " + original_path)
                print("denoised image : " + denoised_path)

                original_image = cv2.imread(original_path, cv2.IMREAD_COLOR)
                denoised_image = cv2.imread(denoised_path, cv2.IMREAD_COLOR)

                # Ensure the images have the same dimensions
                if original_image.shape == denoised_image.shape:
                    psnr = calculate_psnr(original_image, denoised_image)
                    print(f"PSNR for {filename}: {psnr:.2f}")
                    png_name = denoised_path.split("/")[2]
                    noise_name = png_name.split("_")[1] + png_name.split("_")[2] + png_name.split("_")[3]
                    filter_name = png_name.split("_")[9] + png_name.split("_")[10].split(".")[0]
                    string_matrix.append_to_value(noise_name, filter_name, psnr)
                else:
                    print(f"Images {filename} have different dimensions.")

    print("\n\n\n")
    print(string_matrix.display_matrix())
    string_matrix.plot_matrix("plot.png")

    doc = Document()

    doc.add_heading("averages : ")
    for key, value in string_matrix.calculate_average().items():
        print("average of " + key + " : " + str(value))
        doc.add_paragraph("average of " + key + " : " + str(value))

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_heading("std deviations :")

    for key, std_deviation in string_matrix.calculate_std_deviation().items():
        print("std deviation of " + key + " : " + str(std_deviation))
        doc.add_paragraph("std deviation of " + key + " : " + str(std_deviation))

    doc.add_heading("maximum of averages is : "+string_matrix.get_key_with_max_average())
    doc.add_heading("maximum of std deviation is : "+string_matrix.get_key_with_max_std_deviation())

    doc.add_heading("minimum of std deviation is : "+string_matrix.get_key_with_max_std_deviation())
    doc.add_heading("minimum of std deviation is : "+string_matrix.get_key_with_min_std_deviation())
    doc.add_picture("plot.png", width=docx.shared.Inches(5))
    doc.save("document.docx")


if __name__ == "__main__":
    main()
